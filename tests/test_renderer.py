"""DOCX 渲染器的结构性回归测试。"""

from pathlib import Path
import tempfile
import unittest
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt

from resume_generator.config import A4_PAPER, LETTER_PAPER, SUPPORTED_PAPER_SIZES
from resume_generator.models import Contact, Entry, Name, Section
from resume_generator.renderer import render_resume
from resume_generator.styles import (
    BULLET_STYLE,
    ENTRY_TABLE_STYLE,
    REQUIRED_PARAGRAPH_STYLES,
    REQUIRED_TABLE_STYLES,
    SECTION_STYLE,
    SUB_BULLET_STYLE,
    ResumeTheme,
)
from resume_generator.validator import load_json, parse_json, parse_paper_size

FIXTURE = Path(__file__).parent / "fixtures" / "valid_resume.json"


def _xml_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t")))


def _body_blocks(document) -> list[str]:
    blocks: list[str] = []
    for child in document.element.body:
        if child.tag == qn("w:p"):
            blocks.append(_xml_text(child))
        elif child.tag == qn("w:tbl"):
            cells = child.findall(".//w:tr[1]/w:tc", child.nsmap)
            blocks.append(" || ".join(_xml_text(cell) for cell in cells))
    return blocks


class TestRenderResume(unittest.TestCase):
    """使用完整中文 fixture 检查内容、样式和 OOXML 结构。"""

    @classmethod
    def setUpClass(cls) -> None:
        cls._temporary_directory = tempfile.TemporaryDirectory()
        cls.output_path = Path(cls._temporary_directory.name) / "resume.docx"
        data = load_json(FIXTURE)
        name, contacts, sections = parse_json(data)
        sections.extend(
            (
                Section(title="empty section", entries=[]),
                Section(title="blank section", entries=[]),
            )
        )
        render_resume(
            name,
            contacts,
            sections,
            paper_size=parse_paper_size(data),
        ).save(cls.output_path)
        cls.document = Document(cls.output_path)

    @classmethod
    def tearDownClass(cls) -> None:
        cls._temporary_directory.cleanup()

    def test_content_order_and_field_composition(self) -> None:
        self.assertEqual(
            _body_blocks(self.document),
            [
                "王小明",
                "xiaoming.wang@example.com | 138-0000-0000",
                "Education",
                "示例大学 | 金融学学士 || 成都 | 2021.09 - 2025.06",
                "主修公司金融、投资学与计量经济学。",
                "Experience",
                "示例资产管理公司 | 研究实习生 || 上海 | 2025.07 - 至今",
                "协助整理行业研究资料。",
                "维护估值模型中的公开市场数据。",
                "Skills",
                "Python",
                "Excel",
                "SQL",
            ],
        )

    def test_contacts_use_labels_and_explicit_href_targets(self) -> None:
        relationships = self.document.part.rels.values()
        targets = {
            relationship.target_ref
            for relationship in relationships
            if relationship.reltype.endswith("/hyperlink")
        }
        self.assertEqual(
            targets,
            {"mailto:xiaoming.wang@example.com", "tel:+8613800000000"},
        )

    def test_hyperlinked_contacts_are_black_and_single_underlined(self) -> None:
        contact_paragraph = self.document.paragraphs[1]._p
        hyperlinks = contact_paragraph.findall(".//w:hyperlink", contact_paragraph.nsmap)

        self.assertEqual(len(hyperlinks), 2)
        for hyperlink in hyperlinks:
            run_properties = hyperlink.find(".//w:rPr", hyperlink.nsmap)
            self.assertIsNotNone(run_properties)
            assert run_properties is not None

            color = run_properties.find("w:color", run_properties.nsmap)
            underline = run_properties.find("w:u", run_properties.nsmap)
            self.assertIsNotNone(color)
            self.assertIsNotNone(underline)
            assert color is not None
            assert underline is not None
            self.assertEqual(color.get(qn("w:val")), "000000")
            self.assertEqual(underline.get(qn("w:val")), "single")

    def test_blank_href_contacts_are_visible_plain_text(self) -> None:
        document = render_resume(
            Name("测试姓名"),
            (Contact("成都", "   "), Contact("无链接文本", None)),
            (),
            paper_size="A4",
        )

        self.assertEqual(
            _body_blocks(document),
            ["测试姓名", "成都 | 无链接文本"],
        )
        self.assertFalse(
            any(
                relationship.reltype.endswith("/hyperlink")
                for relationship in document.part.rels.values()
            )
        )
        contact_paragraph = document.paragraphs[1]._p
        self.assertFalse(contact_paragraph.findall(".//w:u", contact_paragraph.nsmap))

    def test_empty_fields_and_one_sided_dates_avoid_extra_separators(self) -> None:
        sections = [
            Section(
                title="projects",
                entries=[
                    Entry(
                        title=None,
                        position="分析师",
                        location="",
                        start_date=None,
                        end_date="Present",
                        bullets=[],
                    )
                ],
            )
        ]

        document = render_resume(Name("测试姓名"), (), sections, paper_size="A4")

        self.assertEqual(
            _body_blocks(document),
            ["测试姓名", "Projects", "分析师 || Present"],
        )

    def test_personal_metadata_does_not_remain(self) -> None:
        with ZipFile(self.output_path) as archive:
            package_text = "\n".join(
                archive.read(name).decode("utf-8", errors="ignore")
                for name in archive.namelist()
                if name.endswith((".xml", ".rels"))
            )
        for forbidden in (
            "周一博",
            "18803166626",
            "Eric-Zhou-0302",
            "auto-financial-report-generator-for-Alphapai-LanBaoShu",
        ):
            with self.subTest(forbidden=forbidden):
                self.assertNotIn(forbidden, package_text)

        properties = self.document.core_properties
        self.assertEqual(properties.author, "")
        self.assertEqual(properties.last_modified_by, "")
        self.assertEqual(properties.title, "")
        self.assertEqual(properties.created.year, 2000)
        self.assertEqual(properties.modified.year, 2000)

    def test_required_styles_and_font_slots_exist(self) -> None:
        theme = ResumeTheme()
        expected_style_ids = {
            "Resume Name": "ResumeName",
            "Resume Contact Information": "ResumeContactInformation",
            "Resume Section Heading": "ResumeSectionHeading",
            "Resume Entry Heading": "ResumeEntryHeading",
            "Resume Entry Metadata": "ResumeEntryMetadata",
            "Resume Bullet": "ResumeBullet",
            "Resume Sub Bullet": "ResumeSubBullet",
        }
        for style_name in REQUIRED_PARAGRAPH_STYLES:
            with self.subTest(style=style_name):
                style = self.document.styles[style_name]
                self.assertEqual(style.style_id, expected_style_ids[style_name])
                run_fonts = style.element.rPr.rFonts
                self.assertEqual(run_fonts.get(qn("w:ascii")), theme.en_font)
                self.assertEqual(run_fonts.get(qn("w:hAnsi")), theme.en_font)
                self.assertEqual(run_fonts.get(qn("w:eastAsia")), theme.cn_font)

        for style_name in REQUIRED_TABLE_STYLES:
            style = self.document.styles[style_name]
            self.assertEqual(style.type, WD_STYLE_TYPE.TABLE)
            self.assertEqual(style.style_id, "ResumeEntryTable")

        self.assertEqual(self.document.styles[BULLET_STYLE].font.size, Pt(11))
        self.assertEqual(self.document.styles[SUB_BULLET_STYLE].font.size, Pt(10.5))

    def test_section_heading_uses_title_case_and_true_small_caps(self) -> None:
        section_style = self.document.styles[SECTION_STYLE]
        self.assertIsNotNone(section_style.element.rPr.find(qn("w:smallCaps")))
        section_texts = [
            paragraph.text
            for paragraph in self.document.paragraphs
            if paragraph.style.name == SECTION_STYLE
        ]
        self.assertEqual(section_texts, ["Education", "Experience", "Skills"])

        uppercase_document = render_resume(
            Name("Example Name"),
            (),
            (
                Section("WORK EXPERIENCE", [Entry(bullets=["Example bullet"])]),
                Section("AI PROJECTS", [Entry(bullets=["Example bullet"])]),
                Section("CFA RESEARCH", [Entry(bullets=["Example bullet"])]),
            ),
            paper_size="Letter",
        )
        uppercase_section_texts = [
            paragraph.text
            for paragraph in uppercase_document.paragraphs
            if paragraph.style.name == SECTION_STYLE
        ]
        self.assertEqual(
            uppercase_section_texts,
            ["Work Experience", "Ai Projects", "Cfa Research"],
        )

    def test_real_numbering_and_section_border_are_connected(self) -> None:
        self.assertEqual(
            self.document.styles[BULLET_STYLE].element.pPr.numPr.numId.val,
            1,
        )
        self.assertEqual(
            self.document.styles[SUB_BULLET_STYLE].element.pPr.numPr.numId.val,
            3,
        )
        numbering = self.document.part.numbering_part.element
        num_to_abstract = {
            int(element.get(qn("w:numId"))): element.find(
                qn("w:abstractNumId")
            ).get(qn("w:val"))
            for element in numbering.findall(qn("w:num"))
        }
        abstracts = {
            element.get(qn("w:abstractNumId")): element
            for element in numbering.findall(qn("w:abstractNum"))
        }
        for num_id, style_name in (
            (1, BULLET_STYLE),
            (3, SUB_BULLET_STYLE),
        ):
            with self.subTest(num_id=num_id, style_name=style_name):
                abstract = abstracts[num_to_abstract[num_id]]
                number_format = abstract.find(".//w:numFmt", abstract.nsmap)
                paragraph_style = abstract.find(".//w:pStyle", abstract.nsmap)
                self.assertEqual(number_format.get(qn("w:val")), "bullet")
                self.assertEqual(
                    paragraph_style.get(qn("w:val")),
                    self.document.styles[style_name].style_id,
                )
        section_border = self.document.styles[SECTION_STYLE].element.find(
            "w:pPr/w:pBdr/w:bottom",
            self.document.styles[SECTION_STYLE].element.nsmap,
        )
        self.assertIsNotNone(section_border)
        self.assertEqual(section_border.get(qn("w:val")), "single")

    def test_a4_page_setup_and_half_inch_margins(self) -> None:
        section = self.document.sections[0]
        self.assertEqual(section.orientation, WD_ORIENT.PORTRAIT)
        self.assertAlmostEqual(section.page_width, Mm(210), delta=1000)
        self.assertAlmostEqual(section.page_height, Mm(297), delta=1000)
        for margin in (
            section.top_margin,
            section.bottom_margin,
            section.left_margin,
            section.right_margin,
        ):
            self.assertAlmostEqual(margin, Inches(0.5), delta=1000)

    def test_paper_size_controls_a4_or_letter_page_size(self) -> None:
        expected_dimensions = {
            A4_PAPER: (Mm(210), Mm(297)),
            LETTER_PAPER: (Inches(8.5), Inches(11)),
        }
        for paper_size in SUPPORTED_PAPER_SIZES:
            with self.subTest(paper_size=paper_size):
                document = render_resume(
                    Name("Example Name"),
                    (),
                    (),
                    paper_size=paper_size,
                )
                section = document.sections[0]
                expected_width, expected_height = expected_dimensions[paper_size]
                self.assertAlmostEqual(section.page_width, expected_width, delta=1000)
                self.assertAlmostEqual(
                    section.page_height,
                    expected_height,
                    delta=1000,
                )

    def test_entry_tables_are_fixed_borderless_60_40(self) -> None:
        self.assertEqual(len(self.document.tables), 2)
        table_style = self.document.styles[ENTRY_TABLE_STYLE]
        style_table_properties = table_style.element.find(qn("w:tblPr"))
        layout = style_table_properties.find(qn("w:tblLayout"))
        self.assertEqual(layout.get(qn("w:type")), "fixed")
        borders = style_table_properties.find(qn("w:tblBorders"))
        self.assertTrue(all(edge.get(qn("w:val")) == "nil" for edge in borders))
        margins = style_table_properties.find(qn("w:tblCellMar"))
        self.assertTrue(
            all(edge.get(qn("w:w")) == "0" for edge in margins)
        )
        row_properties = table_style.element.find(qn("w:trPr"))
        self.assertIsNotNone(row_properties.find(qn("w:cantSplit")))
        cell_properties = table_style.element.find(qn("w:tcPr"))
        vertical = cell_properties.find(qn("w:vAlign"))
        self.assertEqual(vertical.get(qn("w:val")), "center")

        for table in self.document.tables:
            with self.subTest(table=_xml_text(table._tbl)):
                self.assertEqual(table.style.name, ENTRY_TABLE_STYLE)
                properties = table._tbl.tblPr
                table_style_ref = properties.find(qn("w:tblStyle"))
                self.assertEqual(
                    table_style_ref.get(qn("w:val")),
                    table_style.style_id,
                )
                self.assertIsNone(properties.find(qn("w:tblLayout")))
                self.assertIsNone(properties.find(qn("w:tblBorders")))
                widths = [
                    int(column.get(qn("w:w")))
                    for column in table._tbl.tblGrid.gridCol_lst
                ]
                self.assertAlmostEqual(widths[0] / sum(widths), 0.6, places=3)
                for cell in table.rows[0].cells:
                    self.assertIsNone(cell._tc.tcPr.find(qn("w:vAlign")))
                    self.assertIsNone(cell._tc.tcPr.find(qn("w:tcMar")))


if __name__ == "__main__":
    unittest.main()
