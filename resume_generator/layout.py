"""简历条目标题与地点日期的双列版式组件。"""

from docx.document import Document as DocumentObject
from docx.table import _Cell, Table

from .helpers import set_run_fonts
from .styles import (
    ENTRY_META_STYLE,
    ENTRY_TITLE_STYLE,
    ResumeEntryTableStyle,
    ResumeTheme,
)


def add_entry_header_table(
    document: DocumentObject,
    left_text: str,
    right_text: str,
    *,
    table_style: ResumeEntryTableStyle,
    theme: ResumeTheme | None = None,
) -> Table:
    """添加固定 60/40、无边框的标题/元信息表格。"""
    active_theme = theme or ResumeTheme()
    table = document.add_table(rows=1, cols=2)
    table_style.apply(table)

    row = table.rows[0]
    _write_cell(
        row.cells[0],
        left_text,
        style_name=ENTRY_TITLE_STYLE,
        theme=active_theme,
    )
    _write_cell(
        row.cells[1],
        right_text,
        style_name=ENTRY_META_STYLE,
        theme=active_theme,
    )
    return table


def _write_cell(
    cell: _Cell,
    text: str,
    *,
    style_name: str,
    theme: ResumeTheme,
) -> None:
    paragraph = cell.paragraphs[0]
    paragraph.style = style_name
    if text:
        run = paragraph.add_run(text)
        set_run_fonts(run, en_font=theme.en_font, cn_font=theme.cn_font)
