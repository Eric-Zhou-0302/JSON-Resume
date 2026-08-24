"""把已校验的数据模型渲染为 Word 文档。"""

import re
from collections.abc import Sequence
from datetime import date

from docx import Document
from docx.document import Document as DocumentObject

from .config import paper_size_for_locale
from .helpers import (
    add_hyperlink,
    clear_core_properties,
    non_empty_text,
    set_run_fonts,
)
from .layout import add_entry_header_table
from .models import Contact, Entry, Name, Section
from .styles import (
    BULLET_STYLE,
    CONTACT_STYLE,
    NAME_STYLE,
    SECTION_STYLE,
    ResumeEntryTableStyle,
    ResumeStyleManager,
    ResumeTheme,
)


def render_resume(
    name: Name,
    contacts: Sequence[Contact],
    sections: Sequence[Section],
    *,
    locale: str,
) -> DocumentObject:
    """从空白 Word 文档构建简历，并注册项目自有样式。"""
    document = Document()
    clear_core_properties(document)

    theme = ResumeTheme()
    paper_size = paper_size_for_locale(locale)
    style_manager = ResumeStyleManager(document, theme, paper_size=paper_size)
    style_manager.apply()
    _render_name(document, name, theme)
    _render_contacts(document, contacts, theme)

    for section in sections:
        visible_entries = [entry for entry in section.entries if _has_content(entry)]
        if not visible_entries:
            continue
        if section.title.strip():
            paragraph = document.add_paragraph(style=SECTION_STYLE)
            run = paragraph.add_run(_format_section_title(section.title))
            set_run_fonts(run, en_font=theme.en_font, cn_font=theme.cn_font)
        for entry in visible_entries:
            _render_entry(
                document,
                entry,
                theme,
                style_manager.entry_table_style,
            )

    return document


def _render_name(
    document: DocumentObject,
    name: Name,
    theme: ResumeTheme,
) -> None:
    text = name.name.strip()
    if not text:
        return
    paragraph = document.add_paragraph(style=NAME_STYLE)
    run = paragraph.add_run(text)
    set_run_fonts(run, en_font=theme.en_font, cn_font=theme.cn_font)


def _render_contacts(
    document: DocumentObject,
    contacts: Sequence[Contact],
    theme: ResumeTheme,
) -> None:
    visible_contacts = [contact for contact in contacts if contact.label.strip()]
    if not visible_contacts:
        return

    paragraph = document.add_paragraph(style=CONTACT_STYLE)
    for index, contact in enumerate(visible_contacts):
        if index:
            separator = paragraph.add_run(" | ")
            set_run_fonts(separator, en_font=theme.en_font, cn_font=theme.cn_font)
        label = contact.label.strip()
        href = contact.href.strip() if contact.href and contact.href.strip() else None
        if href:
            add_hyperlink(
                paragraph,
                label,
                href,
                en_font=theme.en_font,
                cn_font=theme.cn_font,
                color=theme.color,
            )
        else:
            run = paragraph.add_run(label)
            set_run_fonts(run, en_font=theme.en_font, cn_font=theme.cn_font)


def _render_entry(
    document: DocumentObject,
    entry: Entry,
    theme: ResumeTheme,
    table_style: ResumeEntryTableStyle,
) -> None:
    left_text = non_empty_text((entry.title, entry.position), " | ")
    date_text = _format_date_range(entry.start_date, entry.end_date)
    right_text = non_empty_text((entry.location, date_text), " | ")
    if left_text or right_text:
        add_entry_header_table(
            document,
            left_text,
            right_text,
            table_style=table_style,
            theme=theme,
        )

    for bullet in entry.bullets or []:
        text = bullet.strip()
        if not text:
            continue
        paragraph = document.add_paragraph(style=BULLET_STYLE)
        run = paragraph.add_run(text)
        set_run_fonts(run, en_font=theme.en_font, cn_font=theme.cn_font)


def _format_date_range(
    start_date: date | None,
    end_date: date | str | None,
) -> str:
    start_text = start_date.strftime("%Y.%m") if start_date else ""
    if isinstance(end_date, date):
        end_text = end_date.strftime("%Y.%m")
    elif isinstance(end_date, str):
        end_text = end_date.strip()
    else:
        end_text = ""

    if start_text and end_text:
        return f"{start_text} - {end_text}"
    return start_text or end_text


def _format_section_title(title: str) -> str:
    """把英文单词规范为 Title Case，供 Small Caps 显示首字母层级。"""

    def normalize_word(match: re.Match[str]) -> str:
        word = match.group(0)
        return word[:1].upper() + word[1:].lower()

    return re.sub(r"[A-Za-z]+", normalize_word, title.strip())


def _has_content(entry: Entry) -> bool:
    text_fields = (entry.title, entry.position, entry.location)
    if any(value and value.strip() for value in text_fields):
        return True
    if entry.start_date is not None or entry.end_date not in (None, ""):
        return True
    return any(bullet.strip() for bullet in entry.bullets or [])
